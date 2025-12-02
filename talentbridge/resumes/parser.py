import os
import re
import logging
from typing import Dict, List, Optional
import PyPDF2
from docx import Document

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

COMMON_SKILLS = [
    'python', 'javascript', 'java', 'c++', 'c#', 'ruby', 'php', 'swift', 'kotlin', 'go', 'rust',
    'react', 'angular', 'vue', 'node.js', 'express', 'django', 'flask', 'spring', 'laravel',
    'sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch', 'cassandra',
    'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform', 'jenkins', 'ci/cd',
    'machine learning', 'deep learning', 'tensorflow', 'pytorch', 'scikit-learn', 'nlp',
    'data analysis', 'data visualization', 'tableau', 'power bi', 'excel',
    'agile', 'scrum', 'jira', 'git', 'github', 'gitlab',
    'html', 'css', 'sass', 'bootstrap', 'tailwind',
    'rest api', 'graphql', 'microservices', 'api design',
    'linux', 'bash', 'shell scripting', 'networking',
    'project management', 'team leadership', 'communication', 'problem solving',
    'figma', 'sketch', 'adobe xd', 'photoshop', 'illustrator',
    'salesforce', 'sap', 'oracle', 'erp', 'crm',
    'blockchain', 'web3', 'solidity', 'ethereum',
    'mobile development', 'ios', 'android', 'react native', 'flutter',
    'testing', 'selenium', 'cypress', 'jest', 'unit testing', 'qa',
    'security', 'penetration testing', 'owasp', 'cybersecurity',
    'devops', 'sre', 'monitoring', 'logging', 'prometheus', 'grafana'
]

class ResumeParser:
    
    def __init__(self):
        self.skills_pattern = re.compile(r'\b(' + '|'.join(map(re.escape, COMMON_SKILLS)) + r')\b', re.IGNORECASE)
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
        return text
    
    def extract_text_from_docx(self, file_path: str) -> str:
        text = ""
        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
        return text
    
    def extract_text(self, file_path: str) -> str:
        _, ext = os.path.splitext(file_path.lower())
        
        if ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif ext == '.docx':
            return self.extract_text_from_docx(file_path)
        else:
            logger.warning(f"Unsupported file format: {ext}")
            return ""
    
    def extract_skills(self, text: str) -> List[str]:
        text_lower = text.lower()
        found_skills = set()
        
        matches = self.skills_pattern.findall(text_lower)
        found_skills.update([skill.lower() for skill in matches])
        
        return sorted(list(found_skills))
    
    def extract_email(self, text: str) -> Optional[str]:
        email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        match = re.search(email_pattern, text)
        return match.group(0) if match else None
    
    def extract_phone(self, text: str) -> Optional[str]:
        phone_patterns = [
            r'\+?1?[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\+\d{1,3}[-.\s]?\d{10,12}',
            r'\d{10}'
        ]
        for pattern in phone_patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(0)
        return None
    
    def extract_experience_years(self, text: str) -> Optional[int]:
        patterns = [
            r'(\d+)\+?\s*years?\s*(?:of\s*)?(?:experience|exp)',
            r'experience[:\s]*(\d+)\+?\s*years?',
            r'(\d+)\+?\s*years?\s*in\s*(?:the\s*)?(?:industry|field)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                years = int(match.group(1))
                if 0 < years < 50:
                    return years
        return None
    
    def extract_education(self, text: str) -> List[str]:
        education_keywords = [
            r"bachelor'?s?\s*(?:of\s*)?(?:science|arts|engineering|technology)",
            r"master'?s?\s*(?:of\s*)?(?:science|arts|engineering|business|technology)",
            r"ph\.?d\.?",
            r"doctorate",
            r"mba",
            r"b\.?tech",
            r"m\.?tech",
            r"b\.?e\.?",
            r"m\.?e\.?",
            r"b\.?sc",
            r"m\.?sc",
            r"b\.?a\.?",
            r"m\.?a\.?",
            r"bca",
            r"mca",
            r"diploma"
        ]
        
        found_education = []
        for keyword in education_keywords:
            if re.search(keyword, text, re.IGNORECASE):
                match = re.search(keyword, text, re.IGNORECASE)
                if match:
                    found_education.append(match.group(0))
        
        return list(set(found_education))
    
    def parse_resume(self, file_path: str) -> Dict:
        text = self.extract_text(file_path)
        
        if not text:
            return {
                'success': False,
                'error': 'Could not extract text from resume'
            }
        
        skills = self.extract_skills(text)
        email = self.extract_email(text)
        phone = self.extract_phone(text)
        experience_years = self.extract_experience_years(text)
        education = self.extract_education(text)
        
        return {
            'success': True,
            'text': text,
            'skills': skills,
            'email': email,
            'phone': phone,
            'experience_years': experience_years,
            'education': education
        }
